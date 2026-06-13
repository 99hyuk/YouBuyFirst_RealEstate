from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUTPUT = Path("output/submission/너나사_부동산_화면정의서_제출용.docx")
PDF_OUTPUT = Path("output/submission/너나사_부동산_화면정의서_제출용.pdf")

FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
ORANGE = RGBColor(217, 95, 2)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
LIGHT_ORANGE = "FFF3E8"
WHITE = "FFFFFF"

FONT_PATH = Path("C:/Windows/Fonts/malgun.ttf")
FONT_BOLD_PATH = Path("C:/Windows/Fonts/malgunbd.ttf")


def set_run_font(run, size=None, color=None, bold=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def paragraph(text="", size=10.5, color=None, bold=False, after=5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=14, color=BLUE, bold=True)
    else:
        set_run_font(run, size=11.5, color=DARK_BLUE, bold=True)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, size=8.8, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(headers, rows, widths, first_col_orange=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, LIGHT_ORANGE if first_col_orange else LIGHT_GRAY)
        set_cell_margins(cell)
        set_cell_text(cell, header, size=8.5, bold=True, color=ORANGE if first_col_orange else DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            shade_cell(cell, WHITE)
            set_cell_margins(cell)
            set_cell_text(cell, value, size=8.4, bold=(first_col_orange and i == 0), color=ORANGE if first_col_orange and i == 0 else None)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullet(text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(f"- {text}")
    set_run_font(run, size=9.6)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)
section.header_distance = Inches(0.45)
section.footer_distance = Inches(0.45)

styles = doc.styles
styles["Normal"].font.name = FONT
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
styles["Normal"].font.size = Pt(10.5)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = header.add_run("너나사 부동산 화면 정의서")
set_run_font(hr, size=8.5, color=GRAY)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("부동산 관찰형 AI 인사이트 서비스 | 제출용 요약본")
set_run_font(fr, size=8, color=GRAY)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
title_run = title.add_run("너나사 부동산 화면 정의서")
set_run_font(title_run, size=20, color=ORANGE, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(8)
subtitle_run = subtitle.add_run("제출용 요약본 | 기준 화면: https://youbuyfirst-realestate.netlify.app/dashboard | 작성일: 2026-06-04")
set_run_font(subtitle_run, size=9.5, color=GRAY)

paragraph(
    "서비스 요약: 요즘 부동산 관심이 어디로 몰리는지, 사람들의 말과 시장 데이터를 엮어 보여주는 AI 부동산 인사이트 서비스.",
    size=11,
    bold=True,
    color=DARK_BLUE,
    after=5,
)
paragraph(
    "본 서비스는 지역/단지에 대한 커뮤니티 반응, 뉴스/리포트, 실거래·전세·매물·지표, 정책 이벤트를 함께 관찰해 관심 상승 이유와 데이터 근거를 보여준다. "
    "매수·매도·청약·대출 행동을 권유하지 않으며, mock/stale/asOf/provider 상태를 화면에 명확히 표시한다.",
    after=8,
)

heading("1. 화면 구성 요약")
add_table(
    ["화면", "경로", "핵심 목적", "주요 구성"],
    [
        ["대시보드", "/dashboard", "오늘 관심이 몰리는 지역과 이유를 한 화면에서 파악", "부동산 투기 과열 지표, 핵심 지역별 상승률, 지역·단지 반응, 주요 지표, 뉴스/리포트, 데이터 상태"],
        ["지도", "/realestate/map\n/realestate/map/:regionId", "전국·지역별 가격/반응 흐름을 지도 기반으로 탐색", "전국 heat map, 시군구 drilldown, 기간 탭, 상승/하락 색상, 선택 지역 리포트 패널"],
        ["지역 반응", "/realestate/reactions?view=", "커뮤니티에서 급증한 지역·단지와 쟁점을 비교", "언급 급증 카드, 지역 TOP 6, 단지군 TOP 6, 급증 쟁점, 모의 에이전트 근거 로그"],
        ["주요 지표", "/indicators\n/indicators/:category", "부동산 핵심 지표와 지역 반응의 차이를 확인", "가격·거래량, 공급·수급, 수요·심리, 거시·금융 지표, 데이터 신선도, 주요 일정"],
        ["뉴스룸", "/newsroom?feed=&page=", "분석 근거가 되는 뉴스·리포트·영상·원문을 탐색", "뉴스, 정책·통계 리포트, 영상, 블로그/커뮤니티 링크, feed별 목록"],
        ["관심 지역", "/realestate/watchlist", "저장한 지역의 반응 변화와 알림 판단을 추적", "관심 지역 요약, 알림 판단 내역, 원문/공공데이터 후보, 관찰 로그"],
        ["지역/단지 상세", "/realestate/targets/:targetId", "특정 대상의 반응·지표·근거를 리포트 형태로 확인", "한줄 브리핑, 요약 지표, 커뮤니티 반응 추이, 신뢰도, 타임라인, 근거 링크"],
    ],
    [Inches(0.85), Inches(1.35), Inches(1.95), Inches(3.5)],
    first_col_orange=True,
)

heading("2. 핵심 사용자 흐름")
flow_items = [
    "대시보드에서 부동산 투기 과열 지표와 핵심 지역별 상승률을 확인한다.",
    "지도 또는 지역 반응 화면에서 관심이 급증한 지역/단지를 선택한다.",
    "선택한 지역은 지도 리포트 패널 또는 상세 리포트에서 상승·하락 이유와 쟁점을 확인한다.",
    "뉴스룸과 주요 지표 화면에서 원천 데이터, 정책/통계 자료, 커뮤니티 원문 후보를 교차 확인한다.",
    "관심 지역 화면에서 추적 대상과 알림 판단 내역을 관리한다.",
]
for item in flow_items:
    add_bullet(item)

heading("3. ERD 기반 데이터 연결")
add_table(
    ["데이터 영역", "주요 테이블", "사용 화면"],
    [
        ["부동산 대상 정본", "real_estate_targets, real_estate_regions, real_estate_complexes, real_estate_aliases", "전체 화면, 지도, 상세 리포트"],
        ["지도/지역 흐름", "map_boundary_assets, map_features, map_layer_snapshots", "지도, 대시보드 상승률"],
        ["커뮤니티 반응", "crawl_sources, source_boards, community_posts, real_estate_mentions, reaction_analyses, reaction_snapshots, reaction_ranking_rows", "대시보드, 지역 반응, 상세 리포트"],
        ["시장 지표/공공데이터", "real_estate_market_facts, market_indicator_defs, market_indicator_values, market_data_schedules", "대시보드, 주요 지표, 상세 리포트"],
        ["뉴스/정책/근거", "content_items, content_target_links, policy_events, timeline_events, evidence_logs, evidence_log_items", "뉴스룸, 상세 리포트, 근거 로그"],
        ["사용자 관심/알림", "app_users, user_watch_targets, alert_rules, alert_events, observation_logs", "관심 지역, 오른쪽 빠른 패널"],
    ],
    [Inches(1.25), Inches(3.65), Inches(2.65)],
    first_col_orange=True,
)

heading("4. 주요 화면별 데이터 표시 기준")
add_table(
    ["표시 기준", "정의"],
    [
        ["상승/기대", "빨강 계열로 표시하며 가격 상승, 관심 증가, 기대 우세를 의미한다."],
        ["하락/우려", "파랑 계열로 표시하며 가격 하락, 우려 증가, 냉각 신호를 의미한다."],
        ["mock", "실제 API 연결 전 화면 검증용 데이터다."],
        ["stale", "공공데이터 공개 지연 또는 마지막 갱신 기준을 초과한 데이터다."],
        ["asOf/provider", "데이터 기준 시각과 제공 기관을 표시해 실시간처럼 오해하지 않게 한다."],
        ["confidence/sample", "표본 수와 신뢰도를 함께 보여주며, 표본 부족 시 방향 해석을 제한한다."],
    ],
    [Inches(1.35), Inches(6.15)],
)

heading("5. 검사 기준")
checks = [
    "상단 내비게이션에서 모든 주요 화면으로 이동 가능해야 한다.",
    "지도는 전국 지도, 지역 상세 지도, 선택 리포트 패널 상태가 구분되어야 한다.",
    "지역 반응은 종합/순위/반응/근거 로그 탭을 URL query로 식별할 수 있어야 한다.",
    "주요 지표는 가격·거래량, 공급·수급, 수요·심리, 거시·금융 4개 묶음으로 구분되어야 한다.",
    "뉴스룸은 뉴스/리포트/영상/블로그·커뮤니티 feed를 분리해야 한다.",
    "모든 화면은 부동산 자문이 아니라 관찰형 분석 서비스라는 경계를 유지해야 한다.",
    "mock, stale, asOf, provider, confidence 같은 데이터 상태가 숨겨지면 안 된다.",
]
for check in checks:
    add_bullet(check)

heading("6. 구현 우선순위")
add_table(
    ["우선순위", "화면/기능", "이유"],
    [
        ["1", "대시보드", "서비스의 첫인상이며 관심 지역, 반응, 지표, 뉴스가 모두 모이는 중심 화면이다."],
        ["2", "지도", "부동산 서비스의 차별화 지점으로 지역 흐름과 drilldown 경험을 제공한다."],
        ["3", "지역 반응", "커뮤니티 반응 분석 결과가 가장 직접적으로 드러나는 핵심 화면이다."],
        ["4", "주요 지표", "공공데이터와 시장 fact를 연결해 반응 지표의 신뢰도를 보완한다."],
        ["5", "지역/단지 상세", "사용자가 선택한 대상의 이유와 근거를 최종적으로 확인하는 리포트 화면이다."],
        ["6", "뉴스룸/관심 지역", "근거 탐색과 개인 추적 기능으로 확장한다."],
    ],
    [Inches(0.8), Inches(1.6), Inches(5.1)],
)

paragraph("비고: 현재 프론트는 mock/fixture 기반 화면을 포함하며, 실제 구현 시 ERD의 target_id/slug 기준으로 route targetId를 매핑한다.", size=9.2, color=GRAY, after=0)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)


def pdf_paragraph(text, style):
    safe = escape(text).replace("&lt;br/&gt;", "<br/>").replace("\n", "<br/>")
    return Paragraph(safe, style)


def make_pdf_table(headers, rows, col_widths, first_col_orange=False):
    data = [
        [
            pdf_paragraph(
                header,
                pdf_styles["table_header_orange" if first_col_orange else "table_header"],
            )
            for header in headers
        ]
    ]
    for row in rows:
        data.append(
            [
                pdf_paragraph(
                    value,
                    pdf_styles["table_first_col"] if first_col_orange and idx == 0 else pdf_styles["table_cell"],
                )
                for idx, value in enumerate(row)
            ]
        )
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DADCE0")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#FFF3E8" if first_col_orange else "#F2F4F7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def make_pdf():
    pdfmetrics.registerFont(TTFont("MalgunGothic", str(FONT_PATH)))
    pdfmetrics.registerFont(TTFont("MalgunGothicBold", str(FONT_BOLD_PATH)))

    global pdf_styles
    sample = getSampleStyleSheet()
    pdf_styles = {
        "title": ParagraphStyle(
            "title",
            parent=sample["Title"],
            fontName="MalgunGothicBold",
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#D95F02"),
            alignment=TA_LEFT,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            parent=sample["Normal"],
            fontName="MalgunGothic",
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor("#595959"),
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "lead": ParagraphStyle(
            "lead",
            parent=sample["Normal"],
            fontName="MalgunGothicBold",
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "body",
            parent=sample["Normal"],
            fontName="MalgunGothic",
            fontSize=9.3,
            leading=12.5,
            spaceAfter=7,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "h1",
            parent=sample["Heading1"],
            fontName="MalgunGothicBold",
            fontSize=12.5,
            leading=15,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=10,
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "table_header": ParagraphStyle(
            "table_header",
            parent=sample["Normal"],
            fontName="MalgunGothicBold",
            fontSize=7.6,
            leading=9.4,
            textColor=colors.HexColor("#1F4D78"),
            alignment=TA_CENTER,
            wordWrap="CJK",
        ),
        "table_header_orange": ParagraphStyle(
            "table_header_orange",
            parent=sample["Normal"],
            fontName="MalgunGothicBold",
            fontSize=7.6,
            leading=9.4,
            textColor=colors.HexColor("#D95F02"),
            alignment=TA_CENTER,
            wordWrap="CJK",
        ),
        "table_cell": ParagraphStyle(
            "table_cell",
            parent=sample["Normal"],
            fontName="MalgunGothic",
            fontSize=7.2,
            leading=9.3,
            wordWrap="CJK",
        ),
        "table_first_col": ParagraphStyle(
            "table_first_col",
            parent=sample["Normal"],
            fontName="MalgunGothicBold",
            fontSize=7.2,
            leading=9.3,
            textColor=colors.HexColor("#D95F02"),
            wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "bullet",
            parent=sample["Normal"],
            fontName="MalgunGothic",
            fontSize=8.8,
            leading=11.2,
            leftIndent=10,
            bulletIndent=0,
            spaceAfter=2.2,
            wordWrap="CJK",
        ),
        "note": ParagraphStyle(
            "note",
            parent=sample["Normal"],
            fontName="MalgunGothic",
            fontSize=8.2,
            leading=10.5,
            textColor=colors.HexColor("#595959"),
            spaceBefore=4,
            wordWrap="CJK",
        ),
    }

    story = [
        pdf_paragraph("너나사 부동산 화면 정의서", pdf_styles["title"]),
        pdf_paragraph(
            "제출용 요약본 | 기준 화면: https://youbuyfirst-realestate.netlify.app/dashboard | 작성일: 2026-06-04",
            pdf_styles["subtitle"],
        ),
        pdf_paragraph(
            "서비스 요약: 요즘 부동산 관심이 어디로 몰리는지, 사람들의 말과 시장 데이터를 엮어 보여주는 AI 부동산 인사이트 서비스.",
            pdf_styles["lead"],
        ),
        pdf_paragraph(
            "지역/단지에 대한 커뮤니티 반응, 뉴스/리포트, 실거래·전세·매물·지표, 정책 이벤트를 함께 관찰해 관심 상승 이유와 데이터 근거를 보여준다. "
            "매수·매도·청약·대출 행동을 권유하지 않으며, mock/stale/asOf/provider 상태를 명확히 표시한다.",
            pdf_styles["body"],
        ),
        pdf_paragraph("1. 화면 구성 요약", pdf_styles["h1"]),
        make_pdf_table(
            ["화면", "경로", "핵심 목적", "주요 구성"],
            [
                ["대시보드", "/dashboard", "오늘 관심이 몰리는 지역과 이유를 한 화면에서 파악", "과열 지표, 지역별 상승률, 지역·단지 반응, 주요 지표, 뉴스/리포트, 데이터 상태"],
                ["지도", "/realestate/map<br/>/realestate/map/:regionId", "전국·지역별 가격/반응 흐름을 지도 기반으로 탐색", "전국 heat map, 시군구 drilldown, 기간 탭, 선택 지역 리포트 패널"],
                ["지역 반응", "/realestate/reactions?view=", "커뮤니티에서 급증한 지역·단지와 쟁점을 비교", "언급 급증 카드, 지역 TOP 6, 단지군 TOP 6, 급증 쟁점, 근거 로그"],
                ["주요 지표", "/indicators<br/>/indicators/:category", "부동산 핵심 지표와 지역 반응의 차이를 확인", "가격·거래량, 공급·수급, 수요·심리, 거시·금융, 신선도, 일정"],
                ["뉴스룸", "/newsroom?feed=&page=", "분석 근거가 되는 뉴스·리포트·영상·원문을 탐색", "뉴스, 정책·통계 리포트, 영상, 블로그/커뮤니티 링크"],
                ["관심 지역", "/realestate/watchlist", "저장한 지역의 반응 변화와 알림 판단을 추적", "관심 지역 요약, 알림 판단, 원문/공공데이터 후보, 관찰 로그"],
                ["지역/단지 상세", "/realestate/targets/:targetId", "특정 대상의 반응·지표·근거를 리포트 형태로 확인", "한줄 브리핑, 요약 지표, 반응 추이, 신뢰도, 타임라인, 근거 링크"],
            ],
            [22 * mm, 34 * mm, 49 * mm, 75 * mm],
            True,
        ),
        Spacer(1, 5),
        pdf_paragraph("2. 핵심 사용자 흐름", pdf_styles["h1"]),
        *[pdf_paragraph(f"- {item}", pdf_styles["bullet"]) for item in flow_items],
        pdf_paragraph("3. ERD 기반 데이터 연결", pdf_styles["h1"]),
        make_pdf_table(
            ["데이터 영역", "주요 테이블", "사용 화면"],
            [
                ["부동산 대상 정본", "real_estate_targets, real_estate_regions, real_estate_complexes, real_estate_aliases", "전체 화면, 지도, 상세 리포트"],
                ["지도/지역 흐름", "map_boundary_assets, map_features, map_layer_snapshots", "지도, 대시보드 상승률"],
                ["커뮤니티 반응", "crawl_sources, community_posts, real_estate_mentions, reaction_snapshots, reaction_ranking_rows", "대시보드, 지역 반응, 상세 리포트"],
                ["시장 지표/공공데이터", "real_estate_market_facts, market_indicator_defs, market_indicator_values, market_data_schedules", "대시보드, 주요 지표, 상세 리포트"],
                ["뉴스/정책/근거", "content_items, content_target_links, policy_events, timeline_events, evidence_logs", "뉴스룸, 상세 리포트, 근거 로그"],
                ["사용자 관심/알림", "app_users, user_watch_targets, alert_rules, alert_events, observation_logs", "관심 지역, 오른쪽 빠른 패널"],
            ],
            [34 * mm, 104 * mm, 42 * mm],
            True,
        ),
        Spacer(1, 5),
        pdf_paragraph("4. 데이터 표시 기준", pdf_styles["h1"]),
        make_pdf_table(
            ["표시 기준", "정의"],
            [
                ["상승/기대", "빨강 계열로 표시하며 가격 상승, 관심 증가, 기대 우세를 의미한다."],
                ["하락/우려", "파랑 계열로 표시하며 가격 하락, 우려 증가, 냉각 신호를 의미한다."],
                ["mock", "실제 API 연결 전 화면 검증용 데이터다."],
                ["stale", "공공데이터 공개 지연 또는 마지막 갱신 기준을 초과한 데이터다."],
                ["asOf/provider", "데이터 기준 시각과 제공 기관을 표시해 실시간처럼 오해하지 않게 한다."],
                ["confidence/sample", "표본 수와 신뢰도를 함께 보여주며, 표본 부족 시 방향 해석을 제한한다."],
            ],
            [34 * mm, 146 * mm],
        ),
        pdf_paragraph("5. 검사 기준", pdf_styles["h1"]),
        *[pdf_paragraph(f"- {item}", pdf_styles["bullet"]) for item in checks],
        pdf_paragraph("6. 구현 우선순위", pdf_styles["h1"]),
        make_pdf_table(
            ["순위", "화면/기능", "이유"],
            [
                ["1", "대시보드", "관심 지역, 반응, 지표, 뉴스가 모두 모이는 중심 화면"],
                ["2", "지도", "지역 흐름과 drilldown 경험을 제공하는 차별화 화면"],
                ["3", "지역 반응", "커뮤니티 반응 분석 결과가 직접 드러나는 핵심 화면"],
                ["4", "주요 지표", "공공데이터와 시장 fact로 반응 지표의 신뢰도를 보완"],
                ["5", "지역/단지 상세", "선택 대상의 이유와 근거를 최종 확인하는 리포트 화면"],
                ["6", "뉴스룸/관심 지역", "근거 탐색과 개인 추적 기능으로 확장"],
            ],
            [18 * mm, 42 * mm, 120 * mm],
        ),
        pdf_paragraph(
            "비고: 현재 프론트는 mock/fixture 기반 화면을 포함하며, 실제 구현 시 ERD의 target_id/slug 기준으로 route targetId를 매핑한다.",
            pdf_styles["note"],
        ),
    ]

    PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    pdf_doc = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=13 * mm,
        bottomMargin=12 * mm,
        title="너나사 부동산 화면 정의서",
    )
    pdf_doc.build(story)


make_pdf()
print(OUTPUT.resolve())
print(PDF_OUTPUT.resolve())
